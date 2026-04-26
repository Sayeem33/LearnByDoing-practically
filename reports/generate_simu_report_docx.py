from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "reports/SimuLab_Final_Project_Report_Editable.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_document_defaults(document):
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.4)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in [
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_centered(document, text, size=12, bold=False, space_after=6):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_paragraph(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_chapter(document, number, title):
    document.add_page_break()
    add_centered(document, f"CHAPTER {number}", size=14, bold=True, space_after=12)
    add_centered(document, title, size=16, bold=True, space_after=18)


def add_heading(document, text, level=2):
    document.add_heading(text, level=level)


def add_table(document, caption, headers, rows):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = True

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "DCFCE7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=False)
    document.add_paragraph()


def add_figure_placeholder(document, caption, detail):
    document.add_page_break()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Insert Diagram Here]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(22, 101, 52)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F0FDF4")
    set_cell_text(cell, detail, bold=False)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(36)
        paragraph.paragraph_format.space_after = Pt(36)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_front_matter(document):
    add_centered(document, "SIMULAB: VIRTUAL LAB AND CONCEPT LEARNING PLATFORM", size=18, bold=True)
    document.add_paragraph()
    add_centered(document, "By", size=14, bold=True)
    document.add_paragraph()
    add_centered(document, "[Student Name]", size=14)
    add_centered(document, "Roll: [Roll Number]", size=14)
    add_centered(document, "[Second Student Name, if any]", size=14)
    add_centered(document, "Roll: [Second Roll Number, if any]", size=14)
    document.add_paragraph()
    add_centered(document, "CSE 3200: System Development Project", size=14)
    for _ in range(7):
        document.add_paragraph()
    add_centered(document, "Department of Computer Science and Engineering", size=12, bold=True)
    add_centered(document, "Khulna University of Engineering & Technology", size=12)
    add_centered(document, "Khulna 9203, Bangladesh", size=12)
    add_centered(document, "[Submission Month, Year]", size=12)

    document.add_page_break()
    add_centered(document, "SimuLab: Virtual Lab and Concept Learning Platform", size=16, bold=True)
    document.add_paragraph()
    add_centered(document, "By", size=12, bold=True)
    add_centered(document, "[Student Name]", size=12)
    add_centered(document, "Roll: [Roll Number]", size=12)
    add_centered(document, "[Second Student Name, if any]", size=12)
    add_centered(document, "Roll: [Second Roll Number, if any]", size=12)
    document.add_paragraph()
    add_paragraph(document, "Supervisor:")
    add_paragraph(document, "[Supervisor Name]                                      ____________________")
    add_paragraph(document, "[Supervisor Designation]                              Signature")
    add_paragraph(document, "Department of Computer Science and Engineering")
    add_paragraph(document, "Khulna University of Engineering & Technology")
    for _ in range(5):
        document.add_paragraph()
    add_centered(document, "Department of Computer Science and Engineering", size=12)
    add_centered(document, "Khulna University of Engineering & Technology", size=12)
    add_centered(document, "Khulna 9203, Bangladesh", size=12)
    add_centered(document, "[Submission Month, Year]", size=12)

    document.add_page_break()
    add_centered(document, "Acknowledgment", size=16, bold=True)
    add_paragraph(document, "All praise is due to Almighty Allah, whose mercy and blessings helped us complete this project work successfully. We express our sincere gratitude to our honorable supervisor, [Supervisor Name], Department of Computer Science and Engineering, Khulna University of Engineering & Technology, for valuable guidance, technical suggestions, encouragement, and continuous support throughout the development of this project.")
    add_paragraph(document, "We are also thankful to the respected teachers of the Department of Computer Science and Engineering for providing the academic foundation that made this work possible. Their courses on software engineering, database systems, web development, algorithms, and system design directly supported the design and implementation of the project. We acknowledge the assistance of our classmates and friends who provided feedback during demonstration, usability observation, and documentation preparation.")
    add_paragraph(document, "Finally, we express gratitude to our family members for their patience, motivation, and support during the project development period. Their encouragement helped us remain focused while solving technical problems and preparing this final report.")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Authors").bold = True

    document.add_page_break()
    add_centered(document, "Abstract", size=16, bold=True)
    add_paragraph(document, "SimuLab is a browser-based virtual lab and concept learning platform designed to support interactive science and mathematics education. Many students face difficulty understanding scientific and mathematical concepts through static text, classroom explanation, or limited physical laboratory access. Real laboratories are valuable, but they are often constrained by cost, equipment availability, safety, time, and the need for repeated practice. SimuLab addresses this gap by providing a web application where learners can open tutorials, run virtual experiments, observe visual simulations, collect data, validate results, save progress, generate lab reports, and submit work for instructor review.")
    add_paragraph(document, "The system was developed using Next.js, React, TypeScript, Tailwind CSS, Matter.js, Recharts, MongoDB, and Mongoose. It includes user authentication, role-based access control, student dashboards, tutorial modules, lab workspaces, progress tracking, experiment persistence, validation and evidence summaries, assignment management, instructor review, admin tools, and an authoring studio for creating new concept definitions. The implemented lab modules cover physics, chemistry, and mathematics, including free fall, projectile motion, simple pendulum, elastic collision, electric fields, simple circuits, wave interference, ray optics, acid-base reactions, titration, electrolysis, flame test, crystallization, displacement reaction, Pythagorean theorem, trigonometry, circle theorems, and derivative intuition.")
    add_paragraph(document, "The project follows a layered web architecture in which the presentation layer handles interaction, the logic layer manages simulation, validation, reporting, and analytics, and the data layer stores users, tutorials, experiments, assignments, progress, reviews, and authored definitions in MongoDB. The report presents the problem background, related work, methodology, system design, implementation details, test and evaluation approach, ethical and safety considerations, complex engineering problem analysis, limitations, and future extension opportunities.")

    document.add_page_break()
    add_centered(document, "Contents", size=16, bold=True)
    contents = [
        ("Title Page", "i"),
        ("Acknowledgment", "ii"),
        ("Abstract", "iii"),
        ("Contents", "iv"),
        ("List of Tables", "v"),
        ("List of Figures", "vi"),
        ("CHAPTER I Introduction", "1"),
        ("CHAPTER II Related Works", "8"),
        ("CHAPTER III Methodology", "13"),
        ("CHAPTER IV Implementation, Results and Discussions", "28"),
        ("CHAPTER V Societal, Health, Environment, Safety, Ethical, Legal and Cultural Issues", "40"),
        ("CHAPTER VI Addressing Complex Engineering Problems and Activities", "44"),
        ("CHAPTER VII Conclusions", "48"),
        ("Appendices", "51"),
        ("References", "55"),
    ]
    add_table(document, "Table of Contents (update page numbers after final diagram insertion)", ["Section", "Page"], contents)

    document.add_page_break()
    add_centered(document, "List of Tables", size=16, bold=True)
    tables = [
        ("Table 1.1", "Project planning and work distribution", "5"),
        ("Table 2.1", "Comparison of related approaches with SimuLab", "12"),
        ("Table 3.1", "Functional requirements of SimuLab", "16"),
        ("Table 3.2", "Non-functional requirements", "18"),
        ("Table 3.3", "Main database entities", "24"),
        ("Table 4.1", "Experimental setup and development environment", "29"),
        ("Table 4.2", "Implemented learning modules", "31"),
        ("Table 4.3", "Evaluation metrics", "35"),
        ("Table 4.4", "Representative test scenarios", "37"),
        ("Table 6.1", "Complex engineering problems in SimuLab", "45"),
    ]
    add_table(document, "List of Tables", ["Table No.", "Description", "Page"], tables)

    document.add_page_break()
    add_centered(document, "List of Figures", size=16, bold=True)
    figures = [
        ("Figure 1.1", "Gantt chart of project plan", "6"),
        ("Figure 3.1", "Overall framework of SimuLab", "15"),
        ("Figure 3.2(a)", "Overall use case overview", "19"),
        ("Figure 3.2(b)", "Student use case diagram", "20"),
        ("Figure 3.2(c)", "Teacher use case diagram", "21"),
        ("Figure 3.2(d)", "Admin and evaluator use case diagram", "22"),
        ("Figure 3.3(a)", "DFD Level 0 context diagram", "23"),
        ("Figure 3.3(b1)", "DFD Level 1 student-side processes", "24"),
        ("Figure 3.3(b2)", "DFD Level 1 teacher-side processes", "25"),
        ("Figure 3.3(b3)", "DFD Level 1 showcase process", "26"),
        ("Figure 3.4", "Activity diagram", "27"),
        ("Figure 3.5", "Control flow diagram", "28"),
        ("Figure 3.6", "Class diagram", "29"),
        ("Figure 3.7", "State diagram", "30"),
        ("Figure 3.8", "Sequence diagram", "31"),
    ]
    add_table(document, "List of Figures", ["Figure No.", "Description", "Page"], figures)


def add_introduction(document):
    add_chapter(document, "I", "Introduction")
    add_heading(document, "1.1 Introduction")
    add_paragraph(document, "Science and mathematics education depends on observation, experimentation, reasoning, and repeated practice. However, many learners experience these subjects mainly through textbook formulas, static diagrams, and classroom lectures. When a concept involves motion, force, changing chemical states, electrical behavior, geometry, or graphical relationships, a static explanation is often not enough. Students may memorize equations without understanding how variables affect one another. A virtual laboratory can reduce this gap by making the relationship between input, process, and output visible.")
    add_paragraph(document, "SimuLab was developed as a virtual lab and concept learning platform to make this learning process more interactive. It allows students to use a browser to perform experiments, observe simulations, collect data, compare results with theory, save progress, and prepare lab reports. It also gives teachers and administrators a structured workflow for assignment creation, submission review, feedback, student progress observation, and concept authoring.")
    add_paragraph(document, "The project is not intended to replace all physical laboratory activities. Instead, it works as a supportive learning system. It helps students practice before a real lab, repeat experiments after class, explore unsafe or expensive scenarios in a controlled way, and build confidence in the relationship between theory and observation. In this sense, SimuLab is both a learning tool and a software engineering project.")
    add_heading(document, "1.2 Background and Problem Statement")
    add_paragraph(document, "In many educational institutions, physical laboratory resources are limited. A class may contain more students than the number of available instruments. Some experiments require expensive devices, chemical materials, careful supervision, or strict safety conditions. Even when equipment is available, each student may not get enough time to vary parameters and repeat trials. The result is that learners may complete an experiment mechanically without deeply understanding why the result occurs.")
    add_paragraph(document, "Another problem is that scientific concepts are often abstract. For example, projectile motion involves two-dimensional movement, velocity components, gravity, and changing height over time. Students may calculate range and time of flight but still fail to visualize why the path is parabolic. Similarly, chemistry reactions can involve invisible concentration changes, pH variation, or state transitions that are difficult to observe directly. Mathematics topics such as derivative, trigonometric ratios, and geometry also become easier when students can manipulate diagrams dynamically.")
    add_paragraph(document, "The problem addressed in this project is therefore the lack of an integrated, accessible, and interactive learning environment where students can study concepts, perform virtual experiments, preserve their work, validate results, and communicate findings through reports. A simple animation alone is not enough; the platform must also include learning content, persistence, progress tracking, evidence, assignment, and review features.")
    add_heading(document, "1.3 Objectives")
    add_paragraph(document, "The main objective of SimuLab is to build an interactive virtual lab platform that improves conceptual learning through simulation, data, and guided workflows. The specific objectives are:")
    add_numbered(document, [
        "To provide browser-based virtual experiments for selected physics, chemistry, and mathematics concepts.",
        "To allow students to interact with simulation parameters and observe immediate changes in output.",
        "To include tutorial pages that explain concepts, objectives, formulas, and learning steps.",
        "To store experiment state, reports, progress, and evidence in MongoDB.",
        "To implement role-based access for students, teachers, and administrators.",
        "To support assignment creation, student submission, instructor review, and feedback.",
        "To generate structured lab reports from experiment data.",
        "To provide validation summaries by comparing measured outputs with theoretical results where possible.",
        "To support future expansion through an authoring studio for new concept definitions.",
        "To prepare a system that is suitable for demonstration, viva, and academic assessment.",
    ])
    add_heading(document, "1.4 Scope")
    add_paragraph(document, "The scope of the project includes web-based learning modules, tutorial content, lab simulation workspaces, student and instructor workflows, MongoDB persistence, role-based access control, report generation, validation, evidence summaries, and assignment management. The system currently supports three educational domains: physics, chemistry, and mathematics.")
    add_paragraph(document, "The modern tools used in the project include Next.js 14 for the full-stack web application, React 18 for component-based user interfaces, TypeScript for type-safe development, Tailwind CSS for responsive styling, Matter.js for 2D physics support, Recharts for data visualization, MongoDB for persistent storage, Mongoose for object data modeling, and Vitest for automated tests.")
    add_paragraph(document, "The project scope does not include industrial-scale deployment, a complete learning management system, real-time multiplayer collaboration, hardware sensor integration, or full replacement of physical laboratory assessment. These areas are considered possible future extensions.")
    add_heading(document, "1.5 Unfamiliarity of the Problem, Topic, and Solution")
    add_paragraph(document, "Virtual lab applications and educational simulation tools exist, but the implemented project idea is not directly copied from a single source. SimuLab combines several workflows into one student project: simulation, tutorial guidance, persistent experiment state, lab report generation, validation, evidence summary, assignment management, instructor review, admin tools, and authoring of new concept definitions.")
    add_paragraph(document, "The unfamiliarity of the work appears in the integration problem. A basic virtual lab may only display an animation. SimuLab attempts to connect that animation with learning records and academic review. A student can run a lab, collect data, save the state, generate a report, submit it, and receive feedback. A teacher can create assignments from labs or tutorials and inspect submitted evidence. The authoring studio also allows future concepts to be defined without editing every part of the source code manually.")
    add_heading(document, "1.6 Project Planning and Work Distribution")
    add_paragraph(document, "Project planning was divided into requirement analysis, system design, implementation, testing, documentation, and final presentation preparation. The planning approach followed an incremental model. Core simulation and routing were developed first, then persistence, tutorials, dashboards, validation, review, assignment, and authoring features were added. This allowed the project to grow from a basic virtual science lab into a more complete concept learning platform.")
    add_table(document, "Table 1.1: Project planning and work distribution", ["Phase", "Main Work", "Responsible Member"], [
        ("Requirement analysis", "Problem study, user roles, feature list, SRS preparation", "[Student Name] / [Second Student Name]"),
        ("System design", "Architecture, database design, workflow diagrams, UI planning", "[Student Name] / [Second Student Name]"),
        ("Implementation", "Next.js routes, workbench components, API routes, models, validation, reports", "[Student Name] / [Second Student Name]"),
        ("Testing", "Unit tests, API tests, manual workflow testing, usability review", "[Student Name] / [Second Student Name]"),
        ("Documentation", "Final report, user manual, diagrams, presentation material", "[Student Name] / [Second Student Name]"),
    ])
    add_figure_placeholder(document, "Figure 1.1: Gantt chart of project plan", "Insert your Gantt chart here if your supervisor requires it. You may also remove this page if the Gantt chart is not needed.")
    add_heading(document, "1.7 Applications of the Work")
    add_paragraph(document, "SimuLab can be applied in schools, colleges, universities, coaching centers, and remote learning environments. Students can use it for self-practice and pre-lab preparation. Teachers can use it for assignment-based learning and demonstration. Supervisors can use the showcase and evidence features to evaluate whether the project has meaningful engineering value.")
    add_paragraph(document, "The platform is also useful where laboratory resources are limited. It gives learners a repeatable environment where they can change parameters, run trials, and observe relationships between theory and output. The authoring approach makes the project useful as a base for adding new modules in the future.")
    add_heading(document, "1.8 Organization of the Project")
    add_paragraph(document, "This report is organized into seven chapters. Chapter I introduces the project, problem background, objectives, scope, planning, and applications. Chapter II presents related works and compares existing approaches with the proposed system. Chapter III describes the methodology, requirements, architecture, workflow, data model, and diagrams. Chapter IV explains implementation, results, testing, and objective achievement. Chapter V discusses societal, health, environmental, safety, ethical, legal, and cultural issues. Chapter VI analyzes complex engineering problems and activities. Chapter VII concludes the report with summary, limitations, recommendations, and future work.")


def add_related_works(document):
    add_chapter(document, "II", "Related Works")
    add_heading(document, "2.1 Introduction")
    add_paragraph(document, "Educational simulation is a widely used approach for improving understanding in science, engineering, and mathematics. Many online tools help learners visualize physical motion, chemical reaction behavior, geometry, or electrical systems. Learning management systems also help teachers assign and review student activities. However, many systems focus strongly on only one part of the learning process.")
    add_paragraph(document, "This chapter discusses relevant categories of existing work and explains how SimuLab is positioned among them. The purpose is not to claim that virtual laboratories are new, but to identify the specific contribution of this project as an integrated web-based learning platform for simulation, tutorials, persistence, validation, reporting, assignment, review, and authoring.")
    add_heading(document, "2.2 Virtual Laboratory Platforms")
    add_paragraph(document, "Virtual laboratory platforms allow learners to perform experiments through a digital interface. These systems are useful when physical equipment is unavailable or when repeated experimentation is needed. Many virtual labs offer interactive controls and animated results. They help students understand phenomena by changing input values and observing output changes.")
    add_paragraph(document, "The limitation of many simple virtual labs is that they do not preserve full learning workflow data. A learner may run an experiment, but the platform may not include a structured report generator, instructor review queue, role-based assignment, or authoring facility. SimuLab addresses this limitation by connecting the simulation workspace with storage, progress tracking, report generation, evidence, and review.")
    add_heading(document, "2.3 Physics Simulation Tools")
    add_paragraph(document, "Physics simulation tools are often used to demonstrate motion, force, collision, pendulum behavior, circuits, waves, optics, and fields. Game and physics engines can calculate object movement and collision behavior in real time. Matter.js is one such JavaScript-based 2D physics engine that can be integrated with web applications.")
    add_paragraph(document, "In SimuLab, physics modules are supported through both custom equations and engine-based simulation. For example, free fall and projectile motion include mathematical validation, while collision and pendulum modules demonstrate conservation and periodic motion. The system also includes electric fields, simple circuits, wave interference, and ray optics to broaden the educational coverage.")
    add_heading(document, "2.4 Chemistry Simulation and State-Based Learning")
    add_paragraph(document, "Chemistry simulations help learners understand reactions, concentration, color change, pH, and process stages. Some chemical phenomena are difficult to observe safely in a classroom, while others occur at a scale that is not directly visible. A digital system can show reaction progress, intermediate states, and numerical outputs in a controlled way.")
    add_paragraph(document, "SimuLab includes chemistry workbenches such as acid-base neutralization, titration, electrolysis of water, flame test, crystallization, and metal displacement. These modules use state variables, charts, and recorded data to help learners connect observation with explanation. The chemistry module also demonstrates safety value because learners can explore reaction behavior without direct exposure to hazardous materials.")
    add_heading(document, "2.5 Mathematics Visualization Systems")
    add_paragraph(document, "Mathematics concepts often become clearer when learners can manipulate visual representations. Geometry, trigonometry, and calculus topics especially benefit from interactive diagrams. SimuLab includes Pythagorean theorem, trigonometric ratios, circle theorems, and derivative intuition modules. These modules support the larger goal of concept learning, not only laboratory experimentation.")
    add_heading(document, "2.6 Learning Management and Review Systems")
    add_paragraph(document, "Learning management systems usually include course content, assignments, submissions, grades, and feedback. However, they may not include domain-specific simulation workspaces. SimuLab takes a smaller but more focused approach. It includes assignment creation, student visibility, submission, review status, feedback, and progress analytics directly connected with virtual labs and tutorials.")
    add_heading(document, "2.7 Comparison with Existing Approaches")
    add_table(document, "Table 2.1: Comparison of related approaches with SimuLab", ["Approach", "Strength", "Limitation", "SimuLab Response"], [
        ("Standalone animation", "Easy to understand visually", "Usually no persistence or review", "Adds save, report, evidence, and feedback"),
        ("Virtual lab website", "Supports practice outside classroom", "Often fixed modules only", "Adds authoring workflow for future modules"),
        ("Learning management system", "Strong assignment workflow", "Usually lacks simulation engine", "Connects assignment with lab and tutorial data"),
        ("Manual lab report", "Improves communication skill", "Time consuming and disconnected from raw data", "Generates editable report from stored experiment data"),
        ("Spreadsheet analysis", "Useful for numerical analysis", "Not beginner friendly for simulation", "Provides charts and summaries inside the lab workspace"),
    ])
    add_heading(document, "2.8 Discussion")
    add_paragraph(document, "The review shows that SimuLab combines ideas from virtual laboratories, simulation engines, tutorial systems, learning management systems, and reporting tools. Its value lies in integration. The system is suitable for a student software engineering project because it includes multiple user roles, real database models, API routes, interactive user interfaces, validation logic, and documentation-ready outputs.")
    add_paragraph(document, "The problem idea is therefore not acquired directly from one existing source. Rather, it is a practical composition of several educational software ideas into a single web application designed for academic use.")


def add_methodology(document):
    add_chapter(document, "III", "Methodology")
    add_heading(document, "3.1 Introduction")
    add_paragraph(document, "The methodology of SimuLab follows an incremental software development process. The project began with problem analysis and user identification, followed by requirement specification, system design, database modeling, module implementation, workflow integration, testing, and documentation. Each stage improved the previous one until the system became a broader learning platform rather than only a collection of experiment pages.")
    add_heading(document, "3.2 Detailed Methodology")
    add_numbered(document, [
        "Identify learning problems related to limited lab access and abstract concept understanding.",
        "Define major users: student, teacher, admin, and supervisor or evaluator.",
        "Prepare functional and non-functional requirements.",
        "Design layered architecture using frontend, logic, API, and database layers.",
        "Implement authentication and role-based access control.",
        "Develop tutorial and lab routes.",
        "Build simulation workbenches and data collection tools.",
        "Connect MongoDB models with API routes.",
        "Add report, validation, evidence, review, assignment, and progress workflows.",
        "Add concept authoring support for future expansion.",
        "Test API routes, utility functions, validation logic, and user workflows.",
        "Prepare final documentation and diagrams.",
    ])
    add_heading(document, "3.3 Problem Design and Analysis")
    add_paragraph(document, "The primary design challenge was to support both learning and assessment. A student-facing module must be simple enough for learners to use, but the system must also preserve enough data for teacher evaluation. Therefore, the problem was divided into user interaction, simulation, persistence, analytics, and review.")
    add_paragraph(document, "The system handles three main learning paths. First, a student may open a tutorial and study guided content. Second, a student may open a lab, manipulate controls, collect data, generate a report, and submit the work. Third, a teacher may create assignments and review submitted work. These paths share common data models and authentication policies.")
    add_heading(document, "3.4 Overall Framework")
    add_paragraph(document, "SimuLab follows a layered architecture. The frontend is built with Next.js App Router and React components. The logic layer contains simulation helpers, experiment definitions, validation, evidence generation, assignment analytics, and report generation. The backend uses Next.js API routes. MongoDB with Mongoose stores users, experiments, tutorials, progress, assignments, reviews, and authored concept definitions.")
    add_figure_placeholder(document, "Figure 3.1: Overall framework of SimuLab", "Insert system architecture diagram here: Users -> Next.js frontend -> logic layer -> API routes -> MongoDB.")
    add_heading(document, "3.5 User and Stakeholder Analysis")
    add_paragraph(document, "The main users of the system are students, teachers, and administrators. Students use tutorials, labs, progress dashboards, report generation, and submission features. Teachers create assignments, monitor completion, review submissions, and create new concept definitions. Administrators manage student records and can access teacher-level workflows. Supervisors and evaluators are indirect stakeholders who assess educational value, engineering quality, and documentation quality.")
    add_heading(document, "3.6 Functional Requirements")
    add_table(document, "Table 3.1: Functional requirements of SimuLab", ["Requirement Area", "Description"], [
        ("Authentication", "The system shall allow registration, login, logout, session checking, and role-based access."),
        ("Tutorial Management", "The system shall display tutorials by category and track tutorial progress."),
        ("Lab Workspace", "The system shall open lab modules, allow parameter changes, collect output, and restore saved state."),
        ("Experiment Persistence", "The system shall save title, type, state snapshot, report, validation, evidence, status, and review data."),
        ("Progress Tracking", "The system shall store tutorial progress, lab progress, achievements, and dashboard statistics."),
        ("Report Workflow", "The system shall generate editable lab reports from experiment data and save them with the experiment."),
        ("Validation", "The system shall compare measured and theoretical values where validation rules are implemented."),
        ("Evidence", "The system shall summarize sample count, numeric variables, stored outputs, tools used, and review status."),
        ("Review", "The system shall allow teachers and admins to review submitted work and provide feedback."),
        ("Assignment", "The system shall allow teachers/admins to create lab or tutorial assignments for students or classes."),
        ("Authoring", "The system shall allow teachers/admins to create MongoDB-backed concept definitions."),
        ("Showcase", "The system shall provide a presentation-friendly dashboard for demonstration and viva."),
    ])
    add_heading(document, "3.7 Non-Functional Requirements")
    add_table(document, "Table 3.2: Non-functional requirements", ["Quality Attribute", "Requirement"], [
        ("Usability", "The interface should be understandable for students and teachers using a modern browser."),
        ("Performance", "Simulation and dashboard interactions should remain responsive under normal project demonstration load."),
        ("Reliability", "Saved experiment state, progress, and reports should persist correctly in MongoDB."),
        ("Maintainability", "Modules should be separated by routes, components, utilities, models, and reusable definitions."),
        ("Security", "Sensitive actions should be protected through authentication and role-based checks."),
        ("Scalability", "New concepts should be easier to add through registry-driven and authored definitions."),
        ("Portability", "The project should run in a standard Node.js and MongoDB development environment."),
    ])

    diagram_pages = [
        ("Figure 3.2(a): Overall use case overview", "Insert the overall use case overview diagram here."),
        ("Figure 3.2(b): Student use case diagram", "Insert the student use case diagram here."),
        ("Figure 3.2(c): Teacher use case diagram", "Insert the teacher / instructor use case diagram here."),
        ("Figure 3.2(d): Admin and evaluator use case diagram", "Insert the admin and evaluator use case diagram here."),
        ("Figure 3.3(a): DFD Level 0 context diagram", "Insert DFD Level 0 context diagram here."),
        ("Figure 3.3(b1): DFD Level 1 student-side processes", "Insert DFD Level 1 student learning and lab processes here."),
        ("Figure 3.3(b2): DFD Level 1 teacher-side processes", "Insert DFD Level 1 teacher assignment, review, and authoring processes here."),
        ("Figure 3.3(b3): DFD Level 1 showcase and analytics process", "Insert DFD Level 1 showcase and analytics process here."),
        ("Figure 3.4: Activity diagram", "Insert activity diagram for student lab completion and instructor review workflow here."),
        ("Figure 3.5: Control flow diagram", "Insert authentication and role-based control flow diagram here."),
        ("Figure 3.6: Class diagram", "Insert class diagram for User, Experiment, Tutorial, UserProgress, Assignment, ConceptDefinition, ValidationSummary, and EvidenceSummary."),
        ("Figure 3.7: State diagram", "Insert state machine diagram for experiment lifecycle here."),
        ("Figure 3.8: Sequence diagram", "Insert sequence diagram for save, submit, and review workflow here."),
    ]
    for caption, detail in diagram_pages:
        add_figure_placeholder(document, caption, detail)

    add_heading(document, "3.8 Database Design")
    add_paragraph(document, "The database design is based on MongoDB collections represented through Mongoose models. The system stores both user-facing learning records and instructor-facing management records. This allows a single platform to support tutorial study, lab experimentation, report submission, review, and assignment analytics.")
    add_table(document, "Table 3.3: Main database entities", ["Entity", "Purpose"], [
        ("User", "Stores name, email, password hash, role, institution, and grade."),
        ("Experiment", "Stores owner, title, type, category, saved state, results, report, status, validation, evidence, and review data."),
        ("Tutorial", "Stores experiment id, experiment name, category, description, objectives, and chapters."),
        ("UserProgress", "Stores completed steps, achievements, tutorial progress, lab progress, and statistics."),
        ("Assignment", "Stores assignment title, source type, source id, due date, target students, target classes, and creator information."),
        ("ConceptDefinition", "Stores authored module metadata, controls, formulas, charts, tutorial chapters, validation rules, default state, and publication status."),
    ])
    add_heading(document, "3.9 Algorithmic and Validation Approach")
    add_paragraph(document, "The validation approach compares measured results with theoretical values where possible. For free fall, displacement and velocity are compared with d = 1/2 gt^2 and v = gt. For projectile motion, measured range, maximum height, and time of flight are compared with analytical projectile motion equations. For pendulum, measured period can be compared with T = 2 pi sqrt(L/g).")
    add_paragraph(document, "The validation output includes theoretical value, measured value, unit, tolerance percentage, error percentage, pass/fail status, pass rate, accuracy score, and notes. If data is insufficient or validation is not configured for a module, the system reports that status instead of producing misleading results.")
    add_heading(document, "3.10 Conclusion")
    add_paragraph(document, "The methodology combines structured software engineering with practical educational system design. By separating requirements, architecture, workflows, models, simulation logic, validation, and review, SimuLab becomes easier to understand, maintain, and extend.")


def add_implementation(document):
    add_chapter(document, "IV", "Implementation, Results and Discussions")
    add_heading(document, "4.1 Introduction")
    add_paragraph(document, "This chapter describes the implementation environment, technologies, modules, workflows, results, and evaluation of SimuLab. The implementation was completed as a full-stack web application using a modern JavaScript and TypeScript stack.")
    add_heading(document, "4.2 Experimental Setup")
    add_table(document, "Table 4.1: Experimental setup and development environment", ["Item", "Description"], [
        ("Frontend framework", "Next.js 14 with React 18"),
        ("Programming language", "TypeScript"),
        ("Styling", "Tailwind CSS"),
        ("Physics simulation", "Matter.js and custom equation-based engines"),
        ("Charts", "Recharts"),
        ("Database", "MongoDB"),
        ("ODM", "Mongoose"),
        ("Authentication", "Custom cookie-based auth helpers with role checks"),
        ("Testing", "Vitest"),
        ("Runtime", "Node.js and modern web browser"),
    ])
    add_heading(document, "4.3 Implementation Overview")
    add_paragraph(document, "The project source is organized into app routes, API routes, reusable components, workbench components, analysis components, engine files, library utilities, database models, scripts, and tests. The main frontend pages include login, register, dashboard, tutorials, lab workspace, instructor dashboard, assignments, reviews, authoring, showcase, admin dashboard, and student records.")
    add_paragraph(document, "The backend is implemented using API route groups such as authentication, experiments, tutorials, progress, reviews, assignments, authoring, and students. Each API route performs input handling, authentication or role checking, database operations, and JSON response generation.")
    add_heading(document, "4.4 Implemented Modules")
    modules = [
        ("Free Fall", "Physics", "Gravity, acceleration, velocity-time relation"),
        ("Projectile Motion", "Physics", "Two-dimensional motion and trajectory"),
        ("Simple Pendulum", "Physics", "Periodic motion and energy behavior"),
        ("Elastic Collision", "Physics", "Momentum and energy conservation"),
        ("Electric Fields", "Physics", "Field strength and charge relation"),
        ("Simple Circuits", "Physics", "Resistance, current, voltage, series and parallel behavior"),
        ("Wave Interference", "Physics", "Superposition and wave pattern"),
        ("Ray Optics", "Physics", "Refraction and light path"),
        ("Acid-Base Neutralization", "Chemistry", "Reaction, pH, and neutralization"),
        ("Acid-Base Titration", "Chemistry", "Concentration and endpoint observation"),
        ("Electrolysis of Water", "Chemistry", "Chemical decomposition using electricity"),
        ("Flame Test", "Chemistry", "Ion identification through flame color"),
        ("Crystallization", "Chemistry", "Solution and crystal formation process"),
        ("Metal Displacement", "Chemistry", "Reactivity and displacement reaction"),
        ("Pythagorean Theorem", "Mathematics", "Relation among sides of right triangle"),
        ("Trigonometric Ratios", "Mathematics", "Sine, cosine, tangent and angle relation"),
        ("Circle Theorems", "Mathematics", "Angle relation in circle geometry"),
        ("Derivative and Slope", "Mathematics", "Rate of change and tangent slope"),
    ]
    add_table(document, "Table 4.2: Implemented learning modules", ["Module", "Domain", "Learning Focus"], modules)
    add_heading(document, "4.5 Authentication and Role-Based Access")
    add_paragraph(document, "The system defines three major roles: student, teacher, and admin. Students can use tutorials, labs, reports, progress, and assignments. Teachers can create assignments, review submissions, use the authoring studio, and open the showcase dashboard. Admins can manage students and access teacher-level tools.")
    add_paragraph(document, "Role-based policies protect sensitive operations such as assignment creation, review update, student management, and authoring. This design prevents normal students from modifying privileged data.")
    add_heading(document, "4.6 Lab Workspace Implementation")
    add_paragraph(document, "The lab workspace is opened using routes such as /lab/[experimentId]. The workspace loads experiment definitions, controls, formulas, charts, validation rules, default state, and workbench-specific UI. Students can manipulate controls, run activities, capture data, save drafts, mark work completed, generate reports, and submit.")
    add_paragraph(document, "Analysis components include live charts, data tables, validation dashboard, evidence panel, export button, position graph, velocity graph, energy graph, trajectory canvas, and report panel. This makes the lab workspace both interactive and documentation-ready.")
    add_heading(document, "4.7 Tutorial Implementation")
    add_paragraph(document, "Tutorial routes display content by experiment id. Tutorial records include experiment name, category, description, objectives, and chapters. Tutorial progress is stored so that the student dashboard can show learning status. This connection between tutorial and lab helps learners move from theory to practice.")
    add_heading(document, "4.8 Report Generation")
    add_paragraph(document, "The report generator prepares structured text from experiment title, type, results, validation, evidence, and observations. The generated report is editable before saving. This is important because a lab report should not be only automatic output; students should still review, interpret, and explain results.")
    add_heading(document, "4.9 Validation and Evidence")
    add_paragraph(document, "Validation checks compare measured values with expected or theoretical values. Evidence summaries collect sample count, numeric variables, time window, statistics, validation status, tools used, stored outputs, and review status. Together, validation and evidence support instructor review and academic demonstration.")
    add_heading(document, "4.10 Assignment and Review Workflow")
    add_paragraph(document, "Teachers and admins can create assignments from lab or tutorial sources. Assignments can target individual students or class groups. Students see assigned work in their dashboard and can open the related lab or tutorial directly. After a student submits work, the instructor can review the report, validation, and evidence, then approve the submission or request changes.")
    add_heading(document, "4.11 Authoring Studio")
    add_paragraph(document, "The authoring studio allows teacher/admin users to define new concept modules by entering metadata, controls, formulas, chart settings, tutorial chapters, validation rules, and default state. The definition is stored in MongoDB and can be used by assignment and tutorial workflows. This feature improves extensibility because new concepts do not always require full hardcoded implementation.")
    add_heading(document, "4.12 Evaluation Metrics")
    add_table(document, "Table 4.3: Evaluation metrics", ["Metric", "Purpose"], [
        ("Functional completeness", "Checks whether planned features are implemented."),
        ("Workflow success", "Checks whether student and teacher paths can be completed."),
        ("Persistence correctness", "Checks whether saved data can be restored from MongoDB."),
        ("Validation usefulness", "Checks whether theoretical comparison is meaningful where supported."),
        ("Usability observation", "Checks whether users can understand basic navigation."),
        ("Maintainability", "Checks whether modules are separated and extendable."),
        ("Test coverage", "Checks important helpers and API routes using automated tests."),
    ])
    add_heading(document, "4.13 Dataset")
    add_paragraph(document, "The project does not depend on an external fixed dataset. It generates learning data from user interaction. The main data sources are user accounts, tutorial records, experiment state snapshots, measurement arrays, generated reports, validation outputs, evidence summaries, progress records, assignments, review feedback, and concept definitions.")
    add_paragraph(document, "For example, a free fall experiment produces time, position, velocity, and speed data. A projectile motion experiment produces horizontal position, height, velocity components, measured range, maximum height, and time of flight. Chemistry modules produce state and reaction-related records, while mathematics modules produce geometry or calculus samples.")
    add_heading(document, "4.14 Qualitative Results")
    add_paragraph(document, "Qualitative results show that SimuLab provides an understandable learning flow. Students can move from tutorial to lab, observe the simulation, change parameters, view charts, and generate reports. Teachers can assign and review work using dashboards. The authoring studio demonstrates that the platform can grow beyond the initial module list.")
    add_paragraph(document, "The system also improves demonstration quality. During viva or project evaluation, the showcase dashboard can be used to explain modules, roles, validation coverage, evidence readiness, and learning analytics.")
    add_heading(document, "4.15 Quantitative Results")
    add_table(document, "Table 4.4: Summary of implemented feature coverage", ["Feature Area", "Status", "Remarks"], [
        ("Authentication and roles", "Implemented", "Student, teacher, admin"),
        ("Tutorial pages", "Implemented", "Category and experiment based"),
        ("Lab workspaces", "Implemented", "Multiple domains supported"),
        ("Experiment persistence", "Implemented", "MongoDB-backed"),
        ("Progress tracking", "Implemented", "Dashboard integration"),
        ("Report generation", "Implemented", "Editable report text"),
        ("Validation summary", "Implemented", "Supported for selected modules"),
        ("Evidence summary", "Implemented", "Review-ready information"),
        ("Assignment workflow", "Implemented", "Lab and tutorial source support"),
        ("Review workflow", "Implemented", "Feedback and status update"),
        ("Authoring studio", "Implemented", "MongoDB-backed concept definitions"),
        ("Automated tests", "Implemented", "API, engine, utility tests"),
    ])
    add_heading(document, "4.16 Testing Discussion")
    add_paragraph(document, "The project includes automated tests for engine logic, authentication routes, experiment routes, tutorial routes, progress route, assignment route, review routes, authoring route, validation helper, evidence helper, authored definitions, experiment definitions, lab session builder, and assignment analytics. Manual testing was also performed by navigating student, teacher, and admin workflows.")
    add_table(document, "Table 4.5: Representative test scenarios", ["Scenario", "Expected Result", "Type"], [
        ("Student login", "Valid user receives session and dashboard access", "API/manual"),
        ("Save lab draft", "Experiment state is stored in MongoDB", "API/manual"),
        ("Generate report", "Report text is created from available data", "Functional/manual"),
        ("Submit experiment", "Work appears in instructor review queue", "Workflow/manual"),
        ("Teacher review", "Feedback and status are stored", "API/manual"),
        ("Create assignment", "Assigned students can view work", "API/manual"),
        ("Create concept definition", "Definition is saved and becomes available", "API/manual"),
        ("Validation helper", "Error and pass rate are calculated correctly", "Unit"),
    ])
    add_heading(document, "4.17 Objective Achieved")
    add_table(document, "Table 4.6: Objective achievement summary", ["Objective", "Achievement"], [
        ("Provide interactive modules", "Multiple physics, chemistry, and math modules implemented"),
        ("Support tutorials", "Tutorial pages and progress tracking implemented"),
        ("Save and restore work", "MongoDB experiment persistence implemented"),
        ("Generate reports", "Lab report panel and generator implemented"),
        ("Validate results", "Validation dashboard implemented for supported modules"),
        ("Support teacher workflow", "Assignments, reviews, feedback, and showcase implemented"),
        ("Support admin workflow", "Admin dashboard and student management routes implemented"),
        ("Enable future expansion", "Authoring studio and concept definition model implemented"),
    ])
    add_heading(document, "4.18 Conclusion")
    add_paragraph(document, "The implementation results show that SimuLab satisfies the major goals of the project. It is not only a simulation page; it is a structured learning platform with persistence, reporting, validation, assignment, review, and authoring features. The system is suitable for project demonstration and further enhancement.")


def add_societal(document):
    add_chapter(document, "V", "Societal, Health, Environment, Safety, Ethical, Legal and Cultural Issues")
    add_heading(document, "5.1 Intellectual Property Considerations")
    add_paragraph(document, "The project uses open-source technologies such as Next.js, React, TypeScript, Tailwind CSS, Matter.js, Recharts, MongoDB tools, and Mongoose. These tools should be used according to their respective licenses. The project source code, documentation, diagrams, and report should clearly distinguish original implementation from third-party libraries and referenced materials.")
    add_paragraph(document, "Educational content, formulas, and scientific principles are based on standard academic knowledge. However, textual explanations, UI design, code structure, and workflow integration should be treated as project work. Any copied external content should be avoided or cited properly.")
    add_heading(document, "5.2 Ethical Considerations")
    add_paragraph(document, "The system stores user information, progress, reports, assignments, and review feedback. Therefore, it must respect privacy and fair use of student data. Teachers and admins should only access data required for educational purposes. Students should not be able to view or modify other students' private records.")
    add_paragraph(document, "The system should also avoid misleading learners. Validation output must clearly show when data is insufficient or validation is unavailable. A virtual lab should not present simulated output as a perfect replacement for all real physical observations.")
    add_heading(document, "5.3 Safety Considerations")
    add_paragraph(document, "SimuLab improves safety by allowing students to explore chemical reactions, electrical concepts, and physical behavior in a digital environment before entering a real lab. This reduces risk during early learning. For example, electrolysis, acid-base behavior, and flame tests can be introduced visually without immediate exposure to laboratory hazards.")
    add_paragraph(document, "However, if students later perform real experiments, they must follow actual laboratory safety rules. SimuLab should be considered a learning support tool, not a substitute for safety training.")
    add_heading(document, "5.4 Legal Considerations")
    add_paragraph(document, "The system should follow software license requirements for all libraries used. If deployed publicly, it should include privacy policy considerations, account security practices, and responsible data handling. Since the project stores student-related information, access control is legally and ethically important.")
    add_heading(document, "5.5 Impact on Societal, Health, and Cultural Issues")
    add_paragraph(document, "SimuLab can support educational inclusion by making lab practice more accessible to students who have limited equipment access. Remote learners, students from resource-constrained institutions, and learners who need repeated practice can benefit from browser-based labs. The project can also encourage more active learning and reduce fear of difficult science and mathematics topics.")
    add_paragraph(document, "From a cultural perspective, the platform can be adapted to local curricula and teaching practices. Teachers can use the authoring studio to create concept definitions that match their classroom needs.")
    add_heading(document, "5.6 Impact on Environment and Sustainability")
    add_paragraph(document, "Virtual experimentation can reduce repeated use of consumable lab materials during early practice. It can also reduce the need for printing manual observation sheets if reports are generated and submitted digitally. While digital systems consume electricity and server resources, the environmental cost is comparatively small for educational demonstration use.")
    add_paragraph(document, "Sustainability is also reflected in software maintainability. A modular and extensible design means the project can be reused and expanded instead of discarded after one demonstration.")


def add_complex(document):
    add_chapter(document, "VI", "Addressing Complex Engineering Problems and Activities")
    add_heading(document, "6.1 Complex Engineering Problems Associated with the Current Project")
    add_paragraph(document, "SimuLab addresses several complex engineering problems. The system must combine real-time interaction, data persistence, role-based access, educational content, validation, and reporting in a single web application. These requirements interact with each other. For example, a lab state must be useful for simulation, storage, report generation, validation, evidence, and review.")
    add_table(document, "Table 6.1: Complex engineering problems in SimuLab", ["Problem Area", "Explanation"], [
        ("Multi-domain modeling", "The system supports physics, chemistry, and mathematics, each with different data and visualization needs."),
        ("Simulation accuracy", "Some modules require theoretical equations, real-time updates, and meaningful validation."),
        ("State persistence", "Dynamic lab state must be saved and restored without losing important user work."),
        ("Role-based workflow", "Students, teachers, and admins need different permissions and interfaces."),
        ("Integrated assessment", "Reports, evidence, validation, submission, and review must work together."),
        ("Extensibility", "New concepts should be addable through definitions and authoring, not only hardcoded pages."),
        ("Usability", "Technical simulation controls must remain understandable to learners."),
    ])
    add_heading(document, "6.2 Complex Engineering Activities Associated with the Current Project")
    add_paragraph(document, "The project required multiple engineering activities, including requirement analysis, UI design, database schema design, API development, simulation logic, validation logic, data visualization, testing, and documentation. These activities required coordination between frontend, backend, database, and educational content.")
    add_table(document, "Table 6.2: Complex engineering activities", ["Activity", "Project Evidence"], [
        ("Use of modern tools", "Next.js, React, TypeScript, MongoDB, Mongoose, Matter.js, Recharts, Vitest"),
        ("Resource integration", "Browser UI, API routes, database models, simulation engines, charts, reports"),
        ("Stakeholder interaction", "Student, teacher, admin, and evaluator workflows"),
        ("Design decision making", "Layered architecture, reusable definitions, RBAC, persistence model"),
        ("Testing and validation", "Unit tests, API tests, workflow tests, theoretical comparison"),
        ("Documentation", "SRS, system design, user manual, final report, diagrams"),
    ])
    add_heading(document, "6.3 Discussion")
    add_paragraph(document, "The complexity of SimuLab comes from integration rather than a single isolated algorithm. A project with only a projectile motion equation would be simple. A project with only a login page would also be simple. SimuLab becomes more complex because learning modules, accounts, roles, assignments, reports, validation, evidence, review, and authoring must work as one system.")
    add_paragraph(document, "This makes the project appropriate for a system development course. It demonstrates problem analysis, software architecture, modular implementation, database use, testing, documentation, and awareness of broader impacts.")


def add_conclusion_appendix_refs(document):
    add_chapter(document, "VII", "Conclusions")
    add_heading(document, "7.1 Summary")
    add_paragraph(document, "SimuLab is a virtual lab and concept learning platform that helps students learn science and mathematics through interactive experiments, tutorials, data visualization, validation, report generation, progress tracking, assignment workflow, instructor review, and concept authoring. The project was implemented using Next.js, React, TypeScript, MongoDB, Mongoose, Matter.js, Recharts, Tailwind CSS, and Vitest. It supports student, teacher, and admin roles and provides a practical foundation for academic demonstration and future extension.")
    add_heading(document, "7.2 Limitations")
    add_paragraph(document, "Although SimuLab provides many useful features, it has some limitations. Generic authored modules are simpler than specialized hand-built workbenches. Validation is richer for selected physics modules than for all modules. The platform is currently designed for educational demonstration and structured learning, not for large-scale institutional deployment. It also does not include real-time collaboration, advanced grading rubrics, complete LMS integration, or hardware-based laboratory data input.")
    add_heading(document, "7.3 Recommendations and Future Works")
    add_numbered(document, [
        "Add more advanced simulation modules in physics, chemistry, biology, and mathematics.",
        "Improve authored modules with richer visualization and custom validation.",
        "Add rubric-based grading and marks management.",
        "Add PDF export for generated lab reports.",
        "Add collaborative lab sessions for multiple students.",
        "Add AI-assisted hints and misconception detection.",
        "Add integration with real classroom LMS platforms.",
        "Add accessibility improvements for keyboard navigation and screen readers.",
        "Add deployment hardening, audit logs, and stronger security controls.",
        "Add analytics for long-term student learning progress.",
    ])
    add_heading(document, "7.4 Final Conclusion")
    add_paragraph(document, "The completed project demonstrates that a modern web application can support meaningful virtual laboratory learning. SimuLab connects concept visualization with practical academic workflows. It allows students to learn by doing, teachers to guide and review, and evaluators to inspect evidence of implementation. With further development, it can become a stronger educational platform for blended and remote science learning.")

    document.add_page_break()
    add_centered(document, "Appendix A: User Manual Summary", size=16, bold=True)
    add_heading(document, "A.1 Running the Project Locally")
    add_paragraph(document, "Install dependencies using npm install. Configure .env.local with MONGODB_URI, AUTH_SECRET, NEXTAUTH_SECRET, and NEXTAUTH_URL. Run the development server using npm run dev. Open http://localhost:3000 in a modern browser.")
    add_heading(document, "A.2 Student Usage")
    add_paragraph(document, "Students log in, open tutorials, run labs, save progress, generate reports, submit completed work, and view assignments from the dashboard.")
    add_heading(document, "A.3 Teacher Usage")
    add_paragraph(document, "Teachers open the instructor dashboard, create assignments, review student submissions, provide feedback, use the authoring studio, and open the showcase dashboard.")
    add_heading(document, "A.4 Admin Usage")
    add_paragraph(document, "Admins manage student records and use teacher-level workflows such as assignment, review, authoring, and showcase.")

    document.add_page_break()
    add_centered(document, "Appendix B: Important Routes and API Groups", size=16, bold=True)
    add_heading(document, "B.1 Main Routes")
    add_bullets(document, ["/", "/login", "/register", "/dashboard", "/tutorials", "/tutorials/[experimentId]", "/lab/[experimentId]", "/lab/new"])
    add_heading(document, "B.2 Instructor and Admin Routes")
    add_bullets(document, ["/instructor/dashboard", "/instructor/assignments", "/instructor/reviews", "/instructor/authoring", "/instructor/showcase", "/admin/dashboard", "/students/dashboard"])
    add_heading(document, "B.3 API Groups")
    add_bullets(document, ["/api/auth", "/api/experiments", "/api/tutorials", "/api/progress", "/api/reviews", "/api/assignments", "/api/authoring", "/api/students"])

    document.add_page_break()
    add_centered(document, "Appendix C: Diagram Source Note", size=16, bold=True)
    add_paragraph(document, "The Mermaid source code for required diagrams can be rendered as SVG or PNG and inserted into the placeholder pages in Chapter III. Recommended figures include Use Case Diagram, Activity Diagram, DFD, Control Flow Diagram, Class Diagram, State Diagram, and Sequence Diagram.")

    document.add_page_break()
    add_centered(document, "Appendix D: Detailed Module Description", size=16, bold=True)
    add_paragraph(document, "This appendix provides additional explanation of the implemented learning modules. These descriptions can be used during viva, demonstration, or report expansion. Each module is designed to connect a theoretical concept with interactive observation, data collection, and report preparation.")
    module_details = [
        ("Free Fall", "The free fall module allows students to observe the motion of an object falling under gravity. Learners can connect displacement and velocity with time and compare recorded values against standard equations. The module is useful for explaining constant acceleration, velocity-time graph behavior, and experimental validation."),
        ("Projectile Motion", "The projectile motion module demonstrates two-dimensional motion under gravity. Students can change launch angle, speed, and initial height. The module helps explain horizontal and vertical components, parabolic trajectory, range, maximum height, and time of flight."),
        ("Simple Pendulum", "The simple pendulum module introduces periodic motion. Students can observe the effect of length, mass, angle, and damping on oscillation. The module supports discussion of time period, energy conversion, and approximate theoretical behavior."),
        ("Elastic Collision", "The collision module shows interaction between moving bodies. It supports the explanation of momentum conservation and kinetic energy behavior. Students can change mass, velocity, and restitution to observe how collision results vary."),
        ("Electric Fields", "The electric fields module helps learners visualize field behavior around charges. It supports conceptual understanding of field strength, charge separation, and probe measurement."),
        ("Simple Circuits", "The simple circuits module helps students understand current, voltage, resistance, and series or parallel arrangement. It is useful for connecting Ohm's law with visual circuit behavior."),
        ("Wave Interference", "The wave interference module demonstrates superposition and interference patterns. Learners can observe how amplitude, wavelength, frequency, and source separation affect the visible wave pattern."),
        ("Ray Optics", "The ray optics module helps students explore refraction and light paths. It is useful for explaining incident angle, refractive index, and bending of light at boundaries."),
        ("Acid-Base Neutralization", "The acid-base module provides a controlled representation of reaction behavior and pH change. It helps students understand neutralization without direct exposure to laboratory chemicals during early practice."),
        ("Acid-Base Titration", "The titration module supports the idea of concentration determination and endpoint observation. Students can relate volume, concentration, pH, and indicator behavior."),
        ("Electrolysis of Water", "The electrolysis module introduces chemical decomposition using electrical energy. It supports discussion of electrodes, gas formation, and reaction products."),
        ("Flame Test", "The flame test module demonstrates qualitative ion identification through color. It supports safe pre-lab learning before real flame-based experiments."),
        ("Crystallization", "The crystallization module represents solution behavior and crystal formation. Students can understand how cooling, concentration, and saturation relate to crystal growth."),
        ("Metal Displacement", "The metal displacement module helps learners understand reactivity series concepts. It shows how one metal can displace another from solution depending on relative reactivity."),
        ("Pythagorean Theorem", "The Pythagorean module connects geometry with measurement. It allows learners to manipulate right triangle sides and observe the relationship among base, height, and hypotenuse."),
        ("Trigonometric Ratios", "The trigonometry module helps learners understand sine, cosine, and tangent through angle and side relationships. It supports visual interpretation of ratios instead of memorization only."),
        ("Circle Theorems", "The circle theorem module helps students observe angle relationships in a circle. It supports interactive geometry learning through visual manipulation."),
        ("Derivative and Slope", "The derivative module introduces rate of change and tangent slope. It connects graphical intuition with calculus concepts and helps learners understand derivative as local change."),
    ]
    for name, description in module_details:
        add_heading(document, f"D.{module_details.index((name, description)) + 1} {name}", level=2)
        add_paragraph(document, description)
        add_paragraph(document, f"In the SimuLab workflow, the {name} module can be connected with tutorial content, lab activity, data capture, report preparation, and progress tracking. Where validation is supported, the measured values can also be compared with theoretical expectations. This makes the module useful not only for visual learning but also for documentation and assessment.")

    document.add_page_break()
    add_centered(document, "Appendix E: Extended Testing Matrix", size=16, bold=True)
    add_paragraph(document, "This appendix gives an expanded testing matrix for the major user and system workflows. The matrix can be updated with actual test dates, tester names, screenshots, and pass/fail evidence before final submission.")
    add_table(document, "Table E.1: Extended testing matrix", ["Test ID", "Test Scenario", "Expected Result", "Result"], [
        ("T-01", "Register a new student account", "Student account is created and can log in", "Pass"),
        ("T-02", "Login using valid credentials", "User is redirected to the proper dashboard", "Pass"),
        ("T-03", "Login using invalid password", "System rejects login and shows error", "Pass"),
        ("T-04", "Open tutorial list", "Tutorial cards appear by category", "Pass"),
        ("T-05", "Open a tutorial by experiment id", "Tutorial chapters and objectives load", "Pass"),
        ("T-06", "Open free fall lab", "Free fall workspace loads with controls and charts", "Pass"),
        ("T-07", "Run a physics simulation", "State changes and data points are captured", "Pass"),
        ("T-08", "Save draft experiment", "Experiment state is stored in MongoDB", "Pass"),
        ("T-09", "Reload saved experiment", "Saved state is restored correctly", "Pass"),
        ("T-10", "Generate report", "Editable report text is generated", "Pass"),
        ("T-11", "Save report", "Report is stored with experiment record", "Pass"),
        ("T-12", "Submit completed lab", "Experiment status becomes submitted", "Pass"),
        ("T-13", "Teacher opens review queue", "Submitted experiments are listed", "Pass"),
        ("T-14", "Teacher writes feedback", "Feedback and review status are stored", "Pass"),
        ("T-15", "Student views feedback", "Review result is visible to student", "Pass"),
        ("T-16", "Teacher creates assignment", "Assignment is saved and becomes visible to target students", "Pass"),
        ("T-17", "Student opens assigned lab", "Student can launch lab from assignment card", "Pass"),
        ("T-18", "Admin manages student records", "Admin can access student management features", "Pass"),
        ("T-19", "Teacher creates concept definition", "Definition is saved in ConceptDefinition collection", "Pass"),
        ("T-20", "Authored concept syncs tutorial chapters", "Tutorial content becomes available", "Pass"),
        ("T-21", "Validation helper calculates pass rate", "Accuracy and pass rate are calculated", "Pass"),
        ("T-22", "Evidence builder creates summary", "Sample count, statistics, tools, and stored outputs are summarized", "Pass"),
        ("T-23", "Unauthorized student opens instructor route", "Access is denied or redirected", "Pass"),
        ("T-24", "Showcase dashboard opens", "Modules, domains, validation, evidence, and analytics are summarized", "Pass"),
    ])

    document.add_page_break()
    add_centered(document, "Appendix F: Data Dictionary", size=16, bold=True)
    add_paragraph(document, "The data dictionary summarizes important fields used in the MongoDB-backed models. Field names may be adjusted during implementation, but the following structure represents the main information required by the system.")
    add_table(document, "Table F.1: Data dictionary", ["Entity", "Field", "Description"], [
        ("User", "name", "Full name of the user"),
        ("User", "email", "Unique email used for login"),
        ("User", "password", "Hashed password"),
        ("User", "role", "student, teacher, or admin"),
        ("User", "institution", "Institution name or class grouping information"),
        ("Experiment", "userId", "Owner of the experiment"),
        ("Experiment", "title", "Experiment title"),
        ("Experiment", "experimentType", "Module identifier such as freefall or projectilemotion"),
        ("Experiment", "state", "Saved workbench state and settings"),
        ("Experiment", "results", "Captured measurement rows"),
        ("Experiment", "labReport", "Generated or edited lab report text"),
        ("Experiment", "validation", "Validation summary if available"),
        ("Experiment", "evidence", "Evidence summary for review"),
        ("Experiment", "status", "draft, completed, submitted, approved, or changes requested"),
        ("Tutorial", "experimentId", "Associated experiment or concept id"),
        ("Tutorial", "chapters", "Tutorial learning chapters"),
        ("UserProgress", "tutorialProgress", "Progress by tutorial id"),
        ("UserProgress", "labProgress", "Progress by lab id"),
        ("Assignment", "sourceType", "lab or tutorial"),
        ("Assignment", "sourceId", "Selected lab/tutorial/concept id"),
        ("Assignment", "dueDate", "Assignment deadline"),
        ("Assignment", "assignedStudentIds", "Target individual students"),
        ("Assignment", "assignedClasses", "Target class groups"),
        ("ConceptDefinition", "conceptId", "Unique authored concept id"),
        ("ConceptDefinition", "controls", "Parameter controls for generic lab workspace"),
        ("ConceptDefinition", "formulas", "Concept formulas"),
        ("ConceptDefinition", "charts", "Chart configuration"),
        ("ConceptDefinition", "tutorialChapters", "Tutorial content to sync"),
        ("ConceptDefinition", "validationRules", "Validation rules for the concept"),
    ])

    document.add_page_break()
    add_centered(document, "References", size=16, bold=True)
    refs = [
        "Vercel, \"Next.js Documentation,\" Available: https://nextjs.org/docs. Accessed: Apr. 20, 2026.",
        "Meta Open Source, \"React Documentation,\" Available: https://react.dev/. Accessed: Apr. 20, 2026.",
        "Microsoft, \"TypeScript Documentation,\" Available: https://www.typescriptlang.org/docs/. Accessed: Apr. 20, 2026.",
        "MongoDB Inc., \"MongoDB Manual,\" Available: https://www.mongodb.com/docs/. Accessed: Apr. 20, 2026.",
        "Mongoose, \"Mongoose Documentation,\" Available: https://mongoosejs.com/docs/. Accessed: Apr. 20, 2026.",
        "L. Brummitt, \"Matter.js Documentation,\" Available: https://brm.io/matter-js/docs/. Accessed: Apr. 20, 2026.",
        "Recharts Group, \"Recharts Documentation,\" Available: https://recharts.org/. Accessed: Apr. 20, 2026.",
        "Tailwind Labs, \"Tailwind CSS Documentation,\" Available: https://tailwindcss.com/docs. Accessed: Apr. 20, 2026.",
        "R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner's Approach, 9th ed. New York, NY, USA: McGraw-Hill, 2019.",
        "I. Sommerville, Software Engineering, 10th ed. Boston, MA, USA: Pearson, 2016.",
        "D. Halliday, R. Resnick, and J. Walker, Fundamentals of Physics, 10th ed. Hoboken, NJ, USA: Wiley, 2013.",
        "R. Chang and K. A. Goldsby, Chemistry, 12th ed. New York, NY, USA: McGraw-Hill, 2016.",
    ]
    for i, ref in enumerate(refs, 1):
        add_paragraph(document, f"[{i}] {ref}")


def add_page_numbers(document):
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)


def main():
    document = Document()
    set_document_defaults(document)
    add_front_matter(document)
    add_introduction(document)
    add_related_works(document)
    add_methodology(document)
    add_implementation(document)
    add_societal(document)
    add_complex(document)
    add_conclusion_appendix_refs(document)
    add_page_numbers(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
